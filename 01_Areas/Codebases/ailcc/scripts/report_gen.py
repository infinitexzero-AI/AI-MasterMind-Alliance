import os
import sys
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

def markdown_to_pdf(content, output_path):
    """
    Primitive MD to PDF converter using reportlab.
    (In a more advanced setup, we'd use a parser, but for now, we'll write lines).
    """
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter
    y_position = height - 50
    
    for line in content.split('\n'):
        if y_position < 50:
            c.showPage()
            y_position = height - 50
        c.drawString(50, y_position, line)
        y_position -= 15
    
    c.save()
    print(f"✅ PDF Generated: {output_path}")

def markdown_to_docx(content, output_path):
    """
    Converts Markdown content into a Word .docx file.
    """
    doc = Document()
    for line in content.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        else:
            doc.add_paragraph(line)
    
    doc.save(output_path)
    print(f"✅ Word Doc Generated: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python3 report_gen.py <mode: pdf|docx> <input_txt_or_md> <output_path>")
        sys.exit(1)
        
    mode = sys.argv[1]
    input_file = sys.argv[2]
    output_path = sys.argv[3]
    
    if os.path.exists(input_file):
        with open(input_file, 'r') as f:
            content = f.read()
            
        if mode == 'pdf':
            markdown_to_pdf(content, output_path)
        elif mode == 'docx':
            markdown_to_docx(content, output_path)
        else:
            print(f"❌ Unknown mode: {mode}")
    else:
        print(f"❌ Input file not found: {input_file}")
