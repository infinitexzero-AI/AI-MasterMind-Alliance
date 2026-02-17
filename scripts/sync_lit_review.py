import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import os
import re

def add_page_number(run):
    # This adds a basic page number field in a run (compatible with Word)
    # Using a simpler approach for now: standard text is usually fine if we can't inject fields easily
    # But for true APA, page numbers are required.
    pass

def apply_formatting(paragraph, text):
    """Parses simple markdown (bold, italics) and applies to docx paragraph."""
    # Corrected regex: escape both asterisks to treat them as literals, not quantifiers.
    # Pattern for bold (**text**) and italic (*text*)
    # Use capturing group to keep the delimiters in the split results
    parts = re.split(r'(\*{2}.*?\*{2}|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

def update_docx(source_md, target_docx):
    # Create fresh document
    doc = docx.Document()
    
    # --- GLOBAL STYLE SETUP ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)
    
    # --- HEADER (Page Number) ---
    # Simplified: APA 7 Student paper only needs page number in top right
    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # We can't easily add a dynamic "Field" via basic python-docx, 
    # but we can add a placeholder or a simple 1 if it's a short doc.
    # For now, we'll leave it as a placeholder for manual update or just omit to avoid broken fields.
    
    # --- TITLE PAGE (APA 7 Student Paper) ---
    for _ in range(3):
        doc.add_paragraph()
        
    title_text = "Foundations of Health Inquiry: Bridging the Rural Mental Health Divide in New Brunswick"
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.bold = True
    
    doc.add_paragraph() # Extra space
    
    author = doc.add_paragraph("Joel Palk-Ricard")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    affil = doc.add_paragraph("Department of Health Studies, Mount Allison University")
    affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    course = doc.add_paragraph("HLTH 1011: Foundations of Health Inquiry")
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    instr = doc.add_paragraph("Professor San Patten")
    instr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date = doc.add_paragraph("February 17, 2026")
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # --- CONTENT LOADING ---
    with open(source_md, 'r') as f:
        content = f.read()

    # Split by double newlines for paragraphs
    paragraphs_md = content.split('\n\n')
    
    in_references = False

    for p_md in paragraphs_md:
        p_md = p_md.strip()
        if not p_md:
            continue
            
        if p_md.startswith('# '):
            # Title Heading
            p = doc.add_heading('', level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(p_md[2:])
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = docx.shared.RGBColor(0,0,0)
            continue
            
        if p_md.startswith('### '):
            # Level 2 Heading
            p = doc.add_heading('', level=2)
            run = p.add_run(p_md[4:])
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = docx.shared.RGBColor(0,0,0)
            continue

        if 'References' in p_md:
            in_references = True
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("References")
            run.bold = True
            continue

        if p_md.startswith('---'):
            continue

        # Handle body paragraphs vs references
        p = doc.add_paragraph()
        if in_references:
            # Hanging Indent (0.5 inch)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = -Inches(0.5)
            apply_formatting(p, p_md)
        else:
            # Standard first-line indent for body (0.5 inch)
            p.paragraph_format.first_line_indent = Inches(0.5)
            apply_formatting(p, p_md)

    doc.save(target_docx)
    print(f"Successfully updated {target_docx} with high-fidelity APA 7 styling.")

if __name__ == "__main__":
    src = "/Users/infinite27/AILCC_PRIME/AI-MasterMind-Alliance/02_Resources/Academics/HLTH-1011/Refined_HLTH1011_Lit_Review.md"
    tgt = "/Users/infinite27/AILCC_PRIME/AI-MasterMind-Alliance/02_Resources/Academics/HLTH-1011/Palk-Ricard_HLTH1011_Final_Draft.docx"
    update_docx(src, tgt)
