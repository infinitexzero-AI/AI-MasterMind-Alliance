import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_assignment_docx(md_path, docx_path):
    doc = docx.Document()
    
    # Set default style (12pt Times New Roman, Double Spaced)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 2.0
    
    # Configure Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Metadata
    core_props = doc.core_properties
    core_props.author = "Joel Palk-Ricard"
    core_props.title = "GENS2101 Assignment 1 Final"

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    is_title_page = True
    in_references = False

    for i, line in enumerate(lines):
        line = line.strip()
        
        if not line:
            if not is_title_page:
                doc.add_paragraph()
            continue

        # Handle Title Page
        if is_title_page:
            if line.startswith('## '):
                # We hit a body header, so title page is over
                doc.add_page_break()
                is_title_page = False
                # Do NOT continue, let the body logic below process this line
            elif line.startswith('# '):
                # Level 1 Title
                for _ in range(3): doc.add_paragraph()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line[2:])
                run.bold = True
                continue
            elif line.startswith('**'):
                # Title page metadata
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(line.replace('**', ''))
                
                # Check for "Professor" as the last line of title page
                if "Professor" in line:
                    doc.add_page_break()
                    is_title_page = False
                continue
            elif len(line) > 50:
                # Catch-all: if it looks like body text, exit title page mode
                doc.add_page_break()
                is_title_page = False
            else:
                # Probably other metadata or blank lines, skip if in title page mode
                continue

        # Handle Body Content (once is_title_page is False)
        if not is_title_page:
            # Handle Body Sections
            if line.startswith('## '):
                if "References" in line:
                    doc.add_page_break()
                    in_references = True
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run("References")
                    run.bold = True
                    continue
                
                p = doc.add_paragraph()
                run = p.add_run(line[3:])
                run.bold = True
                continue

            # Strip bullet points if in references
            if in_references and line.startswith('* '):
                line = line[2:]

            # Normal Paragraphs
            p = doc.add_paragraph()
            if in_references:
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)
            
            # Advanced Formatting: Bold (**) and Italics (*)
            # Process bold first, then italics in the resulting chunks
            bold_parts = line.split('**')
            for j, b_part in enumerate(bold_parts):
                # Now check for italics in this part
                italic_parts = b_part.split('*')
                for k, i_part in enumerate(italic_parts):
                    if not i_part: continue
                    run = p.add_run(i_part)
                    if j % 2 == 1: run.bold = True
                    if k % 2 == 1: run.italic = True
        
    doc.save(docx_path)
    print(f"File saved to: {docx_path}")

if __name__ == "__main__":
    ACADEMIC_DIR = "/Users/infinite27/AILCC_PRIME/AI-MasterMind-Alliance/02_Resources/Academics/GENS-2101"
    
    # 1. Vision Assignment
    MD_VISION = os.path.join(ACADEMIC_DIR, "Assignment_1_Vision_Refined.md")
    DOCX_VISION = os.path.join(ACADEMIC_DIR, "Joel_Palk-Ricard_GENS2101_Vision_FINAL.docx")
    print(f"Generating Vision Document...")
    generate_assignment_docx(MD_VISION, DOCX_VISION)
    
    # 2. Project Part 1
    MD_PROJECT = os.path.join(ACADEMIC_DIR, "Personal_Project_Part_1_Refined.md")
    DOCX_PROJECT = os.path.join(ACADEMIC_DIR, "Joel_Palk-Ricard_GENS2101_Project_Part1_FINAL.docx")
    print(f"Generating Project Part 1 Document...")
    generate_assignment_docx(MD_PROJECT, DOCX_PROJECT)
