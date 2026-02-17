import os
import re
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Path Configuration
BASE_PATH = "/Users/infinite27/AILCC_PRIME"
ACADEMIC_BASE = os.path.join(BASE_PATH, "02_Resources/Academics")

def extract_frontmatter(md_content):
    """
    Simulated frontmatter extraction from the top of the MD file.
    Looks for Name, Course, Date patterns.
    """
    metadata = {
        "title": "Academic Submission",
        "name": "[User Name]",
        "course": "Natural Resources Management (GENS-2101)",
        "date": "February 17, 2026"
    }
    
    title_match = re.search(r'^# (.*)', md_content, re.MULTILINE)
    if title_match: metadata["title"] = title_match.group(1).strip()
    
    name_match = re.search(r'\*\*Name\*\*:\s*(.*)', md_content)
    if name_match: metadata["name"] = name_match.group(1).strip()
    
    course_match = re.search(r'\*\*Course\*\*:\s*(.*)', md_content)
    if course_match: metadata["course"] = course_match.group(1).strip()

    date_match = re.search(r'\*\*Submission Date\*\*:\s*(.*)', md_content)
    if date_match: metadata["date"] = date_match.group(1).strip()
    
    return metadata

def foundry_convert(md_path, docx_path):
    if not os.path.exists(md_path): return False

    with open(md_path, 'r') as f:
        content = f.read()
    
    meta = extract_frontmatter(content)
    doc = Document()
    
    # Global Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # --- Title Page ---
    doc.add_paragraph("\n" * 5)
    t_p = doc.add_paragraph(meta["title"])
    t_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = t_p.runs[0]
    t.bold = True
    t.font.size = Pt(16)
    
    doc.add_paragraph("\n" * 2)
    for detail in [meta["name"], meta["course"], meta["date"]]:
        p = doc.add_paragraph(detail)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # --- Body ---
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line or line.startswith('# ') or '**Name**' in line or '**Course**' in line or '**Submission Date**' in line:
            continue
        
        # Headers
        if line.startswith('## '):
            p = doc.add_paragraph(line[3:])
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:])
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(13)
        
        # Lists
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p.paragraph_format.line_spacing = 1.5
        
        # Bolding support
        elif '**' in line:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Inches(0.5)
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
        
        # Standard Paragraph
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Inches(0.5)

    doc.save(docx_path)
    return True

if __name__ == "__main__":
    # Test on Assignment 1
    target_md = "/Users/infinite27/AILCC_PRIME/AI-MasterMind-Alliance/02_Resources/Academics/GENS-2101/Assignment_1_Vision.md"
    target_docx = "/Users/infinite27/AILCC_PRIME/AI-MasterMind-Alliance/02_Resources/Academics/GENS-2101/Palk-Ricard_GENS2101_Assignment_1.docx"
    if foundry_convert(target_md, target_docx):
        print(f"✅ Foundry success: {target_docx}")
