import os
import sys
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_academic_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    # APA-ish Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    with open(md_path, 'r') as f:
        lines = f.readlines()

    # --- Title Page (Parsed from headers) ---
    title = "Academic Submission"
    name = "[User Name]"
    date = "February 17, 2026"
    
    for line in lines:
        if line.startswith('# '): title = line[2:].strip()
        if line.startswith('**Name**:'): name = line[9:].strip()
        if line.startswith('**Submission Date**:'): date = line[20:].strip()

    # Add Title Page
    doc.add_paragraph("\n" * 5)
    t_p = doc.add_paragraph(title)
    t_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_p.runs[0].bold = True
    t_p.runs[0].font.size = Pt(16)
    
    doc.add_paragraph("\n" * 2)
    n_p = doc.add_paragraph(name)
    n_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    c_p = doc.add_paragraph("GENS/GENV 2101: Natural Resources Management")
    c_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    d_p = doc.add_paragraph(date)
    d_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # --- Body Content ---
    for line in lines:
        line = line.strip()
        if not line or line.startswith('**Name**') or line.startswith('**Course**') or line.startswith('**Submission Date**') or line.startswith('# '):
            continue
        
        # Headers
        if line.startswith('### '):
            p = doc.add_paragraph(line[4:])
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(13)
        elif line.startswith('## '):
            p = doc.add_paragraph(line[3:])
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(14)
        
        # List Items
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p.paragraph_format.line_spacing = 1.5
        
        # Bolding support
        elif '**' in line:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        # Normal Paragraphs
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Inches(0.5)

    doc.save(docx_path)
    print(f"✅ Success: {docx_path} generated.")

if __name__ == "__main__":
    MD_INPUT = "/Users/infinite27/AILCC_PRIME/AI-MasterMind-Alliance/02_Resources/Academics/GENS-2101/Assignment_1_Vision.md"
    DOCX_OUTPUT = "/Users/infinite27/AILCC_PRIME/AI-MasterMind-Alliance/02_Resources/Academics/GENS-2101/Palk-Ricard_GENS2101_Assignment_1.docx"
    create_academic_docx(MD_INPUT, DOCX_OUTPUT)
